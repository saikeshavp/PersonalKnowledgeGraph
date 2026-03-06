"""
Document Processor Service
Handles text extraction from PDF, DOCX, TXT files and text chunking.
"""

import os
import io
import re
from typing import List, Dict, Any, Optional
from pathlib import Path
import uuid


class DocumentProcessor:
    """Process uploaded documents - extract text and create chunks"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize document processor.
        
        Args:
            chunk_size: Target size for each chunk (in characters)
            chunk_overlap: Overlap between consecutive chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.upload_dir = Path("./uploads")
        self.upload_dir.mkdir(exist_ok=True)
    
    def save_file(self, file_content: bytes, filename: str) -> str:
        """Save uploaded file and return doc_id"""
        doc_id = str(uuid.uuid4())
        file_path = self.upload_dir / f"{doc_id}_{filename}"
        
        with open(file_path, "wb") as f:
            f.write(file_content)
        
        return doc_id, str(file_path)
    
    def extract_text(self, file_path: str, file_type: str) -> str:
        """
        Extract text content from document.
        
        Args:
            file_path: Path to the file
            file_type: File extension (pdf, docx, txt)
        
        Returns:
            Extracted text content
        """
        file_type = file_type.lower().strip(".")
        
        if file_type == "pdf":
            return self._extract_from_pdf(file_path)
        elif file_type == "docx":
            return self._extract_from_docx(file_path)
        elif file_type == "txt":
            return self._extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(file_path)
            text_content = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)
            
            return "\n\n".join(text_content)
        except ImportError:
            raise ImportError("PyPDF2 not installed. Run: pip install pypdf2")
        except Exception as e:
            raise Exception(f"Error extracting PDF: {e}")
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n\n".join(text_content)
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        except Exception as e:
            raise Exception(f"Error extracting DOCX: {e}")
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read()
    
    def chunk_text(
        self, 
        text: str, 
        doc_id: str,
        filename: str
    ) -> List[Dict[str, Any]]:
        """
        Split text into overlapping chunks.
        
        Args:
            text: Full text content
            doc_id: Document ID
            filename: Original filename
        
        Returns:
            List of chunk dictionaries
        """
        # Clean and normalize text
        text = text.strip()
        if not text:
            return []
        
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(text):
            # Find end of chunk
            end = start + self.chunk_size
            
            # Try to break at sentence or paragraph boundary
            if end < len(text):
                # Look for paragraph break
                para_break = text.rfind("\n\n", start, end)
                if para_break > start + self.chunk_size // 2:
                    end = para_break + 2
                else:
                    # Look for sentence break
                    for sep in [". ", "! ", "? ", ".\n", "!\n", "?\n"]:
                        sent_break = text.rfind(sep, start, end)
                        if sent_break > start + self.chunk_size // 2:
                            end = sent_break + len(sep)
                            break
            else:
                end = len(text)
            
            # Extract chunk text
            chunk_text = text[start:end].strip()
            
            if chunk_text:
                chunk_id = f"{doc_id}_chunk_{chunk_index}"
                chunks.append({
                    "id": chunk_id,
                    "text": chunk_text,
                    "doc_id": doc_id,
                    "chunk_index": chunk_index,
                    "source_doc": filename,
                    "start_char": start,
                    "end_char": end
                })
                chunk_index += 1
            
            # Move to next chunk with overlap
            start = end - self.chunk_overlap
            if start >= len(text) - self.chunk_overlap:
                break
        
        return chunks
    
    def process_document(
        self, 
        file_content: bytes, 
        filename: str
    ) -> Dict[str, Any]:
        """
        Complete document processing pipeline.
        
        Args:
            file_content: File bytes
            filename: Original filename
        
        Returns:
            Processing result with doc_id and chunks
        """
        # Determine file type
        file_type = filename.split(".")[-1].lower()
        
        # Save file
        doc_id, file_path = self.save_file(file_content, filename)
        
        # Extract text
        text = self.extract_text(file_path, file_type)
        
        # Preprocess: collapse whitespace, strip control chars
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
        text = re.sub(r'[ \t]+', ' ', text)  # Collapse spaces/tabs but keep newlines
        text = re.sub(r'\n{3,}', '\n\n', text)  # Max 2 consecutive newlines
        
        # Create chunks
        chunks = self.chunk_text(text, doc_id, filename)
        
        return {
            "doc_id": doc_id,
            "filename": filename,
            "file_type": file_type,
            "text_length": len(text),
            "chunks": chunks,
            "chunk_count": len(chunks)
        }


# Singleton instance
document_processor = DocumentProcessor()
